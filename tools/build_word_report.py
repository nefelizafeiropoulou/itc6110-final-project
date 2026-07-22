import re
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
BODY = ROOT / "tools" / "report_body.md"
IMAGE_DIR = ROOT / "tmp" / "notebook"
OUTPUT = ROOT / "output" / "ITC6110_NLP_Project_Report.docx"

NAVY = "17365D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1F2933"
MUTED = "667085"
LIGHT = "E8EEF5"
LIGHTER = "F4F6F9"
WHITE = "FFFFFF"
GOLD = "B0893D"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_font(run, name="Calibri", size=None, bold=None, italic=None, color=INK):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_field(paragraph, instruction, display=""):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = display
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        run._r.append(node)
    set_font(run, size=9, color=MUTED)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(10)
    caption.paragraph_format.keep_with_next = False


def configure_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def add_running_furniture(section):
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("ITC6110 | Natural Language Processing Project")
    set_font(r, size=8.5, bold=True, color=MUTED)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Page ")
    set_font(r, size=9, color=MUTED)
    add_field(p, "PAGE", "1")


def add_cover(doc):
    for _ in range(5):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(18)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("NATURAL LANGUAGE PROCESSING")
    set_font(r, size=11, bold=True, color=GOLD)
    p.paragraph_format.space_after = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Mining the Visitor Experience")
    set_font(r, size=30, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(26)
    r = p.add_run("Sentiment, Topics, Explainability and Retrieval-Augmented Generation\nfrom Rome Colosseum Reviews")
    set_font(r, size=15, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("ITC 6110 - Natural Language Processing")
    set_font(r, size=12, bold=True, color=INK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Final Group Project | Spring Semester 2026")
    set_font(r, size=11, color=MUTED)

    for _ in range(3):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)

    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [3120, 3120, 3120], indent=0)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    entries = [("CORPUS", "8,285 reviews"), ("BEST MACRO-F1", "0.47 | XGBoost"), ("PIPELINE", "NLP to RAG")]
    set_repeat_table_header(table.rows[0])
    for cell, (label, value) in zip(table.rows[0].cells, entries):
        set_cell_shading(cell, LIGHTER)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(label)
        set_font(r, size=8, bold=True, color=GOLD)
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value)
        set_font(r, size=11, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Prepared from the executed ITC6110 project notebook and recorded outputs")
    set_font(r, size=9.5, italic=True, color=MUTED)
    doc.add_page_break()


def add_contents(doc):
    p = doc.add_paragraph("Contents", style="Heading 1")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(14)
    toc = [
        "Abstract",
        "1. Introduction",
        "2. Data Collection and Understanding",
        "3. Data Preprocessing and Normalization",
        "4. Feature Engineering and Text Visualization",
        "5. Supervised Sentiment Classification",
        "6. Explainable AI with SHAP",
        "7. Unsupervised Topic Modelling",
        "8. Retrieval-Augmented Generation",
        "9. Streamlit Interface and Deployment Status",
        "10. Challenges, Limitations, and Reproducibility",
        "11. Future Work",
        "12. Conclusion",
        "References",
        "Appendix A. Recorded Configuration Summary",
    ]
    for item in toc:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.18 if re.match(r"\d+\.", item) else 0)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        set_font(r, size=10.5, bold=item in {"Abstract", "References"} or item.startswith("Appendix"), color=INK)
    doc.add_page_break()


def add_inline_runs(paragraph, text):
    text = text.replace("`", "")
    parts = re.split(r"(\*[^*]+\*)", text)
    for part in parts:
        if part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def add_figure(doc, marker):
    match = re.match(r"\[\[FIGURE:([^|]+)\|(.+)\]\]", marker)
    filename, caption = match.groups()
    path = IMAGE_DIR / filename
    with Image.open(path) as image:
        width, height = image.size
    max_width = 6.15
    max_height = 5.0 if filename == "cell-089-output-1.png" else 6.15
    aspect = width / height
    draw_width = min(max_width, max_height * aspect)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    picture = p.add_run().add_picture(str(path), width=Inches(draw_width))
    picture._inline.docPr.set("title", caption.split(".", 1)[0])
    picture._inline.docPr.set("descr", caption)
    cp = doc.add_paragraph(caption, style="Caption")
    cp.paragraph_format.keep_together = True


def add_markdown_table(doc, rows):
    parsed = [[cell.strip() for cell in row.strip().strip("|").split("|")] for row in rows]
    headers = parsed[0]
    body = parsed[2:]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    if len(headers) == 6:
        widths = [2700, 1332, 1332, 1332, 1332, 1332]
    elif len(headers) == 3:
        widths = [900, 3000, 5460]
    else:
        widths = [9360 // len(headers)] * len(headers)
        widths[-1] += 9360 - sum(widths)
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_font(r, size=9, bold=True, color=NAVY)
    for row_data in body:
        row = table.add_row()
        for idx, text in enumerate(row_data):
            cell = row.cells[idx]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx in (0, 1, 2) and len(headers) == 3 else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text.replace("approximately", "approx."))
            set_font(r, size=8.5, color=INK)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)


def parse_body(doc):
    lines = BODY.read_text(encoding="utf-8").splitlines()
    index = 0
    while index < len(lines):
        line = lines[index].strip()
        if not line:
            index += 1
            continue
        if line.startswith("[[FIGURE:"):
            add_figure(doc, line)
            index += 1
            continue
        if line.startswith("|"):
            rows = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                rows.append(lines[index].strip())
                index += 1
            add_markdown_table(doc, rows)
            continue
        if line.startswith("### "):
            doc.add_paragraph(line[4:], style="Heading 3")
        elif line.startswith("## "):
            doc.add_paragraph(line[3:], style="Heading 2")
        elif line.startswith("# "):
            title = line[2:]
            if title in {"References", "Appendix A. Recorded Configuration Summary"}:
                doc.add_page_break()
            doc.add_paragraph(title, style="Heading 1")
        else:
            p = doc.add_paragraph()
            add_inline_runs(p, line)
            if line.startswith("Blei,") or line.startswith("Chen,") or line.startswith("Hochreiter,") or line.startswith("Lewis,") or line.startswith("Lundberg,") or line.startswith("Mikolov,") or line.startswith("Pedregosa,") or line.startswith("Reimers,") or line.startswith("Rome Colosseum") or line.startswith("van der Maaten"):
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.first_line_indent = Inches(-0.25)
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_after = Pt(6)
        index += 1


def add_document_metadata(doc):
    props = doc.core_properties
    props.title = "Mining the Visitor Experience: NLP Analysis of Rome Colosseum Reviews"
    props.subject = "ITC6110 Natural Language Processing Final Group Project"
    props.author = "ITC6110 Project Team"
    props.keywords = "NLP, sentiment analysis, Word2Vec, XGBoost, BiLSTM, LDA, SHAP, RAG"
    props.comments = "Generated from the executed project notebook and its recorded outputs."


def main():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    for section in doc.sections:
        configure_section(section)
    add_document_metadata(doc)
    add_running_furniture(doc.sections[0])
    add_cover(doc)
    add_contents(doc)
    parse_body(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
